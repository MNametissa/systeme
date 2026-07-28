#!/usr/bin/env python3
"""Vérifie un livrable rendu avant qu'il parte — la porte de sortie.

Usage : verif_livrable.py livrable.docx

BLOQUANT (code 1) : tout résidu de gabarit ({{ }}, {% %}, {marqueur}).
ÉNUMÉRÉ : chaque chiffre du document avec sa position et son contexte —
l'instrument énumère, le modèle (ou l'humain) juge chiffre par chiffre
qu'une source existe ; un chiffre injustifiable = le document ne part pas.
SIGNALÉ : typographie française (espace insécable manquante avant ; ! ? :
et dans les guillemets « ») — décompte et exemples, sans bloquer.

Livrable sans texte = code 2, jamais un OK muet.
"""
import re
import sys
from pathlib import Path

import docx

RESIDU = re.compile(r"\{\{[^}]*\}\}|\{%[^}]*%\}|\{[#/]?[a-z_0-9]+(?:\.[a-z_0-9]+)*\}")
CHIFFRE = re.compile(r"\d+(?:[   .,/]\d+)*(?:\s?%)?")
# insécable (U+00A0 ou U+202F) exigée avant ; ! ? : et à l'intérieur de « »
PONCT_SANS_INSECABLE = re.compile(r"(?<=[^\s  ])[;!?]|(?<= )[;!?:]")
GUILLEMET_SANS_INSECABLE = re.compile(r"«(?![  ])|(?<![  ])»")


def contexte(txt, m, marge=25):
    return txt[max(0, m.start() - marge):m.end() + marge].strip()


def main():
    if len(sys.argv) != 2:
        print(__doc__)
        return 1
    livrable = Path(sys.argv[1])
    doc = docx.Document(str(livrable))
    paragraphes = [(i, p.text) for i, p in enumerate(doc.paragraphs, 1) if p.text.strip()]
    cellules = [(f"t{ti}", c.text) for ti, table in enumerate(doc.tables, 1)
                for row in table.rows for c in row.cells if c.text.strip()]
    blocs = paragraphes + cellules
    if not blocs:
        print(f"ÉCHEC — {livrable.name} ne contient aucun texte : "
              "rien à vérifier n'est pas un succès")
        return 2

    residus = []
    for pos, txt in blocs:
        for m in RESIDU.finditer(txt):
            residus.append((pos, m.group(0)))
    if residus:
        for pos, r in residus[:10]:
            print(f"ÉCHEC — résidu de gabarit p{pos} : {r}")
        return 1

    chiffres = []
    vus = set()
    for pos, txt in blocs:
        for m in CHIFFRE.finditer(txt):
            cle = (pos, m.group(0), contexte(txt, m))
            if cle not in vus:          # les cellules fusionnées se répètent
                vus.add(cle)
                chiffres.append(cle)
    print(f"{len(chiffres)} chiffres à sourcer :")
    for pos, val, ctx in chiffres:
        print(f"  p{pos}: {val!r} — «{ctx}»")

    typo = []
    for pos, txt in blocs:
        for m in PONCT_SANS_INSECABLE.finditer(txt):
            typo.append((pos, contexte(txt, m, 15)))
        for m in GUILLEMET_SANS_INSECABLE.finditer(txt):
            typo.append((pos, contexte(txt, m, 15)))
    print(f"typographie : {len(typo)} insécable(s) manquante(s)")
    for pos, ctx in typo[:8]:
        print(f"  p{pos}: «{ctx}»")

    print(f"\nverdict : zéro résidu ; {len(chiffres)} chiffres à juger un par un")
    return 0


if __name__ == "__main__":
    sys.exit(main())
