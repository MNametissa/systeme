#!/usr/bin/env python3
"""Porte de la machine-redaction — T1 gabarit PTF.

Éprouve instruments/convertir_gabarit.py sur la vraie PTF, puis le rendu
docxtpl sur un contexte à 7 phases (le 6 codé en dur doit déborder).
Sortie non nulle au premier échec. Zéro vérification exécutée = échec.
"""
import json
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

ICI = Path(__file__).resolve().parent
RACINE = ICI.parent          # machine-redaction/
DEPOT = RACINE.parent        # machines/
INSTRUMENT = RACINE / "instruments" / "convertir_gabarit.py"
SOURCE = DEPOT / "Modèles de docx" / "PTF - GINUTECH.docx"
CONTEXTE = ICI / "fixtures" / "contexte_ptf.json"

MOUSTACHE = re.compile(r"\{[#/]?[a-z_0-9]+(?:\.[a-z_0-9]+)*\}")
JINJA = re.compile(r"\{\{|\{%")

verifs = 0
def ok(nom):
    global verifs
    verifs += 1
    print(f"  OK  {nom}")

def echec(nom, detail=""):
    print(f"ÉCHEC {nom}" + (f" — {detail}" if detail else ""))
    sys.exit(1)

def textes_docx(chemin):
    """Concatène les w:t de chaque partie XML du docx."""
    parties = {}
    with zipfile.ZipFile(chemin) as z:
        for n in z.namelist():
            if n.endswith(".xml"):
                xml = z.read(n).decode("utf-8", errors="replace")
                parties[n] = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml))
    return parties

def octets(chemin, membre):
    with zipfile.ZipFile(chemin) as z:
        return z.read(membre)

tmp = Path(tempfile.mkdtemp(prefix="mr-test-"))
gabarit = tmp / "ptf-standard.docx"

# ── 1. Conversion sur la vraie PTF ─────────────────────────────────────────
if not SOURCE.exists():
    echec("source", f"{SOURCE} absent")
r = subprocess.run([sys.executable, str(INSTRUMENT), str(SOURCE), str(gabarit)],
                   capture_output=True, text=True)
if r.returncode != 0:
    echec("conversion", f"code {r.returncode}\n{r.stdout}{r.stderr}")
if not re.search(r"\b\d+ marqueur", r.stdout):
    echec("conversion: décompte", f"pas de décompte dans la sortie:\n{r.stdout}")
ok("conversion: code 0 et décompte annoncé")

for nom, txt in textes_docx(gabarit).items():
    restes = MOUSTACHE.findall(txt)
    if restes:
        echec("conversion: marqueurs moustache restants", f"{nom}: {restes}")
ok("conversion: zéro marqueur {moustache} restant")

# ── 2. Mise en page intacte (styles, en-têtes, pieds, médias) ─────────────
with zipfile.ZipFile(SOURCE) as z:
    a_preserver = [n for n in z.namelist()
                   if n == "word/styles.xml"
                   or re.match(r"word/(header|footer)\d+\.xml$", n)
                   or n.startswith("word/media/")]
for membre in a_preserver:
    if octets(SOURCE, membre) != octets(gabarit, membre):
        echec("préservation", f"{membre} modifié par la conversion")
ok(f"préservation: {len(a_preserver)} parties identiques à l'octet")

# ── 3. Rendu docxtpl, 7 phases ────────────────────────────────────────────
from docxtpl import DocxTemplate
contexte = json.loads(CONTEXTE.read_text())
rendu = tmp / "ptf-rendu.docx"
doc = DocxTemplate(str(gabarit))
doc.render(contexte)
doc.save(str(rendu))
ok("rendu: docxtpl aboutit")

corps = textes_docx(rendu)["word/document.xml"]
residus = JINJA.findall(corps) + MOUSTACHE.findall(corps)
if residus:
    echec("rendu: résidus de gabarit", str(residus[:10]))
ok("rendu: zéro résidu {{ }} {% %} {moustache}")

for attendu in ["ACME SARL", "12 500 000", "Accompagnement",
                "Rapport d'optimisation", "Formateur",
                "Bon pour accord", "Steve FASSEU"]:
    if attendu not in corps:
        echec("rendu: contenu fixture absent", attendu)
if corps.count("ACME SARL") < 3:   # page de garde, tableau infos, signature
    echec("rendu: signature", "le bloc signature ne porte pas le nom du client")
ok("rendu: valeurs du contexte présentes (7e phase, bloc signature)")

# ── 4. Les boucles débordent le 6 codé en dur ─────────────────────────────
if corps.count("Dossier de spécifications") != 1:
    echec("boucle financière", "ligne modèle dupliquée ou absente")
if len(re.findall(r"PHASE 7", corps)) < 1:
    echec("boucle phases", "PHASE 7 absente du rendu")
xml_rendu = zipfile.ZipFile(rendu).read("word/document.xml").decode("utf-8")
i = xml_rendu.find("Phases du projet")
fin = xml_rendu.find("</w:tbl>", i)
lignes_planning = len(re.findall(r"<w:tr[ >]", xml_rendu[xml_rendu.rfind("<w:tbl>", 0, i):fin]))
if lignes_planning != 2 + 7:   # en-tête durée + en-tête semaines + 7 phases
    echec("boucle planning", f"{lignes_planning} lignes au lieu de 9")
ok("boucles: 7 phases rendues (financier, prose, planning)")

tbl_planning = xml_rendu[xml_rendu.rfind("<w:tbl>", 0, i):fin]
vertes = 0
for tr in re.findall(r"<w:tr[ >].*?</w:tr>", tbl_planning, re.S):
    txt_tr = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", tr)).strip()
    if txt_tr.startswith("Phase"):
        n_tr = tr.count('w:fill="0d8a52"')
        vertes += n_tr
        if txt_tr.startswith("Phase 1") and n_tr != 1:
            echec("gantt", f"phase 1 (durée 1) : {n_tr} cellules ombrées au lieu de 1")
if vertes != 9:   # somme des durées de la fixture (1+1+2+2+1+1+1)
    echec("gantt", f"{vertes} cellules ombrées au lieu de 9 — le planning doit "
          "se remplir depuis debut/duree de la base de faits")
ok("gantt: 9 semaines ombrées depuis la base de faits, phase 1 correcte")

# ── 5. Contenu libre : une variable reçoit un sous-document entier ────────
# (le gabarit est une coque, pas un formulaire : des sections nouvelles
#  peuvent naître au rendu, avec les styles nommés du modèle)
from docx.oxml.ns import qn
doc2 = DocxTemplate(str(gabarit))
sd = doc2.new_subdoc()
titre = sd.add_paragraph("Périmètre hors ERP")
pPr = titre._p.get_or_add_pPr()
ps = pPr.makeelement(qn("w:pStyle"), {qn("w:val"): "Heading2"})
pPr.append(ps)
sd.add_paragraph("Cette section n'existe pas dans le modèle : elle est née au rendu.")
t = sd.add_table(rows=2, cols=2)
t.rows[0].cells[0].text = "Sujet"
t.rows[0].cells[1].text = "Décision"
t.rows[1].cells[0].text = "Site vitrine"
t.rows[1].cells[1].text = "Reporté en phase 2"
contexte2 = dict(contexte, solution_description=sd)
rendu2 = tmp / "ptf-rendu-libre.docx"
doc2.render(contexte2)
doc2.save(str(rendu2))
xml2 = zipfile.ZipFile(rendu2).read("word/document.xml").decode("utf-8")
if "Périmètre hors ERP" not in xml2 or "Reporté en phase 2" not in xml2:
    echec("contenu libre", "section née au rendu absente")
i = xml2.find("Périmètre hors ERP")
pPr_seg = xml2[xml2.rfind("<w:p", 0, i):i]
if 'w:pStyle w:val="Heading2"' not in pPr_seg:
    echec("contenu libre", "le titre né au rendu ne porte pas le style Heading2")
ok("contenu libre: section + tableau nés au rendu, styles du modèle hérités")

# ── 6. Gabarit partiel : couverture seule, en-têtes/pieds conservés ───────
DECOUPE = RACINE / "instruments" / "decouper_gabarit.py"
coquille = tmp / "ptf-couverture.docx"
r = subprocess.run([sys.executable, str(DECOUPE), str(gabarit), str(coquille),
                    "--jusqua", "Sommaire"], capture_output=True, text=True)
if r.returncode != 0:
    echec("découpe", f"code {r.returncode}\n{r.stdout}{r.stderr}")
corps_coq = textes_docx(coquille)["word/document.xml"]
for present in ["PROPOSITION", "{{ client_name }}"]:
    if present not in corps_coq:
        echec("découpe: couverture amputée", present)
for absent in ["Sommaire", "Macro planning", "Le prestataire"]:
    if absent in corps_coq:
        echec("découpe: corps non retiré", absent)
ok("découpe: couverture gardée, corps retiré")

for membre in a_preserver:
    if not membre.startswith("word/media/"):
        if octets(SOURCE, membre) != octets(coquille, membre):
            echec("découpe: préservation", f"{membre} modifié")
ok("découpe: styles et en-têtes/pieds identiques à l'octet")

with zipfile.ZipFile(coquille) as z:
    noms = z.namelist()
    rels_refs = set()
    for n in noms:
        if n.endswith(".rels"):
            rels_refs |= set(re.findall(r'Target="media/([^"]+)"', z.read(n).decode()))
    medias = {n.split("/")[-1] for n in noms if n.startswith("word/media/")}
    doc_xml_coq = z.read("word/document.xml").decode()
    ids_utilises = set(re.findall(r'(?:embed|id|link)="(rId\d+)"', doc_xml_coq))
    rels_doc = z.read("word/_rels/document.xml.rels").decode()
    ids_declares = set(re.findall(r'Id="(rId\d+)"[^>]*Type="[^"]*/image"', rels_doc))
if medias - rels_refs:
    echec("découpe: médias orphelins", str(medias - rels_refs))
if ids_utilises & ids_declares != ids_declares & ids_utilises or not ids_declares <= ids_utilises:
    echec("découpe: rels d'images pendantes", str(ids_declares - ids_utilises))
if len(medias) >= 7:
    echec("découpe: propreté", f"{len(medias)} médias conservés sur 7 — rien purgé")
ok(f"découpe: propre ({len(medias)} médias conservés sur 7, zéro orphelin)")

doc3 = DocxTemplate(str(coquille))
doc3.render(contexte)
rendu3 = tmp / "couverture-rendue.docx"
doc3.save(str(rendu3))
txt3 = textes_docx(rendu3)["word/document.xml"]
if "ACME SARL" not in txt3 or JINJA.findall(txt3) or MOUSTACHE.findall(txt3):
    echec("découpe: rendu", "résidus ou contexte absent")
r = subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf",
                    "--outdir", str(tmp), str(rendu3)],
                   capture_output=True, text=True, timeout=120)
r2 = subprocess.run(["pdfinfo", str(tmp / "couverture-rendue.pdf")],
                    capture_output=True, text=True)
pages = int(re.search(r"Pages:\s+(\d+)", r2.stdout).group(1)) if r2.returncode == 0 else -1
if pages < 1 or pages > 2:
    echec("découpe: pdf", f"{pages} pages pour une couverture")
ok(f"découpe: la couverture rendue tient en {pages} page(s), PDF valide")

# ── 7. Refus d'une ancre introuvable : code 2 ─────────────────────────────
r = subprocess.run([sys.executable, str(DECOUPE), str(gabarit), str(tmp / "x.docx"),
                    "--jusqua", "TEXTE INTROUVABLE XYZ"], capture_output=True, text=True)
if r.returncode != 2:
    echec("découpe: ancre introuvable", f"code {r.returncode} au lieu de 2")
ok("découpe: ancre introuvable → code 2, jamais un OK muet")

# ── 8. Marqueurs fragmentés entre plusieurs runs (le cas Word réel) ───────
import docx
frag = tmp / "fragmente.docx"
d = docx.Document()
p = d.add_paragraph()
p.add_run("Le client ")
p.add_run("{cli")           # marqueur coupé en trois runs,
p.add_run("ent_na")         # comme Word le fait après édition
p.add_run("me} est engagé.")
d.save(str(frag))
r = subprocess.run([sys.executable, str(INSTRUMENT), str(frag), str(tmp / "frag-conv.docx")],
                   capture_output=True, text=True)
if r.returncode != 0:
    echec("fragmentation", f"code {r.returncode}\n{r.stdout}{r.stderr}")
txt_frag = textes_docx(tmp / "frag-conv.docx")["word/document.xml"]
if "{{ client_name }}" not in txt_frag or MOUSTACHE.findall(txt_frag):
    echec("fragmentation", f"marqueur coupé non converti : {txt_frag!r}")
ok("fragmentation: marqueur coupé en 3 runs converti")

# ── 9. Les autres modèles balisés se convertissent et se rendent ──────────
AUTRES = ["Approche de gestion.docx", "cahier de conception.docx",
          "Contrat de maintenance.docx", "Contrat partenariat.docx",
          "contrat-sous-traitance.docx", "Modèle CDC.docx"]
for nom_doc in AUTRES:
    src = DEPOT / "Modèles de docx" / nom_doc
    conv = tmp / ("conv-" + nom_doc)
    r = subprocess.run([sys.executable, str(INSTRUMENT), str(src), str(conv)],
                       capture_output=True, text=True)
    if r.returncode != 0:
        echec(f"modèle {nom_doc}", f"conversion code {r.returncode}\n{r.stdout}{r.stderr}")
    for partie, txt in textes_docx(conv).items():
        if MOUSTACHE.findall(txt):
            echec(f"modèle {nom_doc}", f"moustaches restantes dans {partie}")
    with zipfile.ZipFile(src) as z:
        fixes = [n for n in z.namelist()
                 if n == "word/styles.xml" or re.match(r"word/(header|footer)\d+\.xml$", n)
                 or n.startswith("word/media/")]
    for membre in fixes:
        if octets(src, membre) != octets(conv, membre):
            echec(f"modèle {nom_doc}", f"{membre} modifié")
    dtpl = DocxTemplate(str(conv))
    dtpl.render({})          # variables absentes → vide, mais zéro résidu exigé
    vide_rendu = tmp / ("rendu-" + nom_doc)
    dtpl.save(str(vide_rendu))
    txt_r = textes_docx(vide_rendu)["word/document.xml"]
    if JINJA.findall(txt_r) or MOUSTACHE.findall(txt_r):
        echec(f"modèle {nom_doc}", f"résidus au rendu : {(JINJA.findall(txt_r) + MOUSTACHE.findall(txt_r))[:6]}")
    ok(f"modèle {nom_doc}: converti, préservé, rendu sans résidu")

# ── 9bis. CDC : les boucles imbriquées portent le contenu profond ─────────
ctx_cdc = {"functional_modules": [{
    "module_name": "GRC",
    "submodules": [{
        "submodule_name": "Prospection",
        "features": [{
            "feature_name": "Créer un prospect",
            "use_case": {
                "use_case_name": "Création d'une fiche prospect",
                "actors": "Agent commercial",
                "main_scenario_steps": [
                    {"step_number": "1", "step_description": "Ouvrir la fiche prospect"},
                    {"step_number": "2", "step_description": "Saisir les coordonnées"}],
                "alternative_scenarios": [
                    {"alt_step_number": "2a", "alt_description": "Le prospect existe déjà"}]}}]}]}]}
cdc_conv = tmp / "conv-Modèle CDC.docx"
dcdc = DocxTemplate(str(cdc_conv))
dcdc.render(ctx_cdc)
cdc_rendu = tmp / "cdc-profond.docx"
dcdc.save(str(cdc_rendu))
txt_cdc = textes_docx(cdc_rendu)["word/document.xml"]
for attendu in ["Ouvrir la fiche prospect", "Saisir les coordonnées",
                "Le prospect existe déjà", "Agent commercial"]:
    if attendu not in txt_cdc:
        echec("CDC imbriqué", f"contenu profond absent : {attendu}")
ok("CDC: 4 niveaux d'imbrication rendus (module→sous-module→feature→étapes)")

# ── 10. Refus d'une accolade hors convention (docxtpl la ferait exploser) ─
brace = tmp / "accolade.docx"
d = docx.Document()
d.add_paragraph("Un texte avec {marqueur_valide} et une { accolade perdue.")
d.save(str(brace))
r = subprocess.run([sys.executable, str(INSTRUMENT), str(brace), str(tmp / "b.docx")],
                   capture_output=True, text=True)
if r.returncode != 1 or "accolade" not in r.stdout.lower():
    echec("refus accolade", f"code {r.returncode}, sortie : {r.stdout!r}")
ok("refus accolade: { hors convention → code 1, nommée dans la sortie")

# ── 11. Refus du vide : docx sans marqueur → code 2 ───────────────────────
import docx
vide = tmp / "sans-marqueur.docx"
d = docx.Document()
d.add_paragraph("Document sans aucun marqueur.")
d.save(str(vide))
r = subprocess.run([sys.executable, str(INSTRUMENT), str(vide), str(tmp / "out.docx")],
                   capture_output=True, text=True)
if r.returncode != 2:
    echec("refus du vide", f"code {r.returncode} au lieu de 2")
ok("refus du vide: 0 marqueur → code 2")

# ── 12. Baliser : document nu + carte → gabarit moustache ─────────────────
BALISE = RACINE / "instruments" / "baliser_gabarit.py"
nu = tmp / "document-nu.docx"
d = docx.Document()
p = d.add_paragraph()
p.add_run("Contrat avec ")
p.add_run("ACME")            # littéral fragmenté entre runs, comme Word le fait
p.add_run(" SARL, signé le 5 mai. Paiement le 5 mai.")
d.save(str(nu))
carte = tmp / "carte.json"
carte.write_text(json.dumps([
    {"texte": "ACME SARL", "nom": "client_name"},
    {"texte": "5 mai", "nom": "date_paiement", "occurrences": [2]},
]))
r = subprocess.run([sys.executable, str(BALISE), str(nu), str(tmp / "balise.docx"),
                    "--carte", str(carte)], capture_output=True, text=True)
if r.returncode != 0:
    echec("balisage", f"code {r.returncode}\n{r.stdout}{r.stderr}")
txt_b = textes_docx(tmp / "balise.docx")["word/document.xml"]
if "{client_name}" not in txt_b:
    echec("balisage", f"littéral fragmenté non balisé : {txt_b!r}")
if "signé le 5 mai" not in txt_b or "Paiement le {date_paiement}" not in txt_b:
    echec("balisage: occurrences", f"restriction d'occurrence ignorée : {txt_b!r}")
ok("balisage: littéral fragmenté balisé, occurrence ciblée respectée")

carte_vide = tmp / "carte-vide.json"
carte_vide.write_text(json.dumps([{"texte": "TEXTE ABSENT DU DOCUMENT", "nom": "x"}]))
r = subprocess.run([sys.executable, str(BALISE), str(nu), str(tmp / "b2.docx"),
                    "--carte", str(carte_vide)], capture_output=True, text=True)
if r.returncode != 2 or "TEXTE ABSENT" not in r.stdout:
    echec("balisage: refus du vide", f"code {r.returncode}, sortie : {r.stdout!r}")
ok("balisage: entrée sans capture → code 2, nommée dans la sortie")

# ── 13. Chaîne complète sur le vrai Contrat de prestation ─────────────────
CARTE_CP = RACINE / "cartes" / "contrat-prestation.json"
src_cp = DEPOT / "Modèles de docx" / "Contrat de prestation.docx"
r = subprocess.run([sys.executable, str(BALISE), str(src_cp), str(tmp / "cp-balise.docx"),
                    "--carte", str(CARTE_CP)], capture_output=True, text=True)
if r.returncode != 0:
    echec("prestation: balisage", f"code {r.returncode}\n{r.stdout}{r.stderr}")
r = subprocess.run([sys.executable, str(INSTRUMENT), str(tmp / "cp-balise.docx"),
                    str(tmp / "cp-gabarit.docx")], capture_output=True, text=True)
if r.returncode != 0:
    echec("prestation: conversion", f"code {r.returncode}\n{r.stdout}{r.stderr}")
ctx_cp = {"contract_number": "CTR00099", "contract_date": "02/09/2026",
          "client_name": "M. ONANA", "client_company_name": "NKOULOU SARL",
          "project_name": "Refonte du portail RH", "total_amount_num": "1 200 000",
          "total_amount_text": "Un million deux cent mille"}
dcp = DocxTemplate(str(tmp / "cp-gabarit.docx"))
dcp.render(ctx_cp)
dcp.save(str(tmp / "cp-rendu.docx"))
txt_cp = textes_docx(tmp / "cp-rendu.docx")["word/document.xml"]
for attendu in ["CTR00099", "NKOULOU SARL", "Refonte du portail RH", "1 200 000"]:
    if attendu not in txt_cp:
        echec("prestation: rendu", f"valeur absente : {attendu}")
for interdit in ["CTR00005", "M. AZEBAZE", "750 000", "Globalys"]:
    if interdit in txt_cp:
        echec("prestation: rendu", f"donnée du contrat d'origine encore présente : {interdit}")
ok("prestation: chaîne carte → balisage → conversion → rendu, données d'origine purgées")

# ── 15. Le contrat gabarit ↔ contexte : variables énumérées et vérifiées ──
VARIABLES = RACINE / "instruments" / "variables_gabarit.py"
r = subprocess.run([sys.executable, str(VARIABLES), str(gabarit)],
                   capture_output=True, text=True)
if r.returncode != 0:
    echec("variables", f"code {r.returncode}\n{r.stdout}{r.stderr}")
if "15 variables" not in r.stdout or "mission_summary" not in r.stdout:
    echec("variables: énumération", r.stdout)
ok("variables: 15 variables du gabarit PTF énumérées (mission_summary comprise)")

ctx_ampute = dict(contexte)
del ctx_ampute["date"]
ctx_file = tmp / "ctx-ampute.json"
ctx_file.write_text(json.dumps(ctx_ampute))
r = subprocess.run([sys.executable, str(VARIABLES), str(gabarit),
                    "--contexte", str(ctx_file)], capture_output=True, text=True)
if r.returncode != 1 or "date" not in r.stdout:
    echec("variables: contexte incomplet", f"code {r.returncode}\n{r.stdout}")
ctx_file.write_text(json.dumps(contexte))
r = subprocess.run([sys.executable, str(VARIABLES), str(gabarit),
                    "--contexte", str(ctx_file)], capture_output=True, text=True)
if r.returncode != 0:
    echec("variables: contexte complet refusé", r.stdout)
ok("variables: variable manquante nommée (code 1), contexte complet accepté")

r = subprocess.run([sys.executable, str(VARIABLES), str(vide)],
                   capture_output=True, text=True)
if r.returncode != 2:
    echec("variables: refus du vide", f"code {r.returncode}")
ok("variables: gabarit sans variable → code 2")

# ── 16. Remplir : strict, sans résidu, sans écrasement ────────────────────
REMPLIR = RACINE / "instruments" / "remplir_gabarit.py"
livrable = tmp / "livrable-ptf.docx"
r = subprocess.run([sys.executable, str(REMPLIR), str(gabarit), str(ctx_file),
                    str(livrable)], capture_output=True, text=True)
if r.returncode != 0 or not livrable.exists():
    echec("remplir", f"code {r.returncode}\n{r.stdout}{r.stderr}")
if "ACME SARL" not in textes_docx(livrable)["word/document.xml"]:
    echec("remplir", "contexte absent du livrable")
pdf_livrable = livrable.with_suffix(".pdf")
if not pdf_livrable.exists() or pdf_livrable.stat().st_size < 10000:
    echec("remplir: pdf", "le livrable doit sortir en docx ET en pdf")
ok("remplir: livrable produit en rendu strict, docx et pdf")

reglages = octets(livrable, "word/settings.xml").decode("utf-8")
if "w:updateFields" not in reglages:
    echec("remplir: sommaire", "updateFields absent du docx — Word n'actualisera pas les champs")
r = subprocess.run(["pdftotext", "-f", "2", "-l", "2", str(pdf_livrable), "-"],
                   capture_output=True, text=True)
if "Développements Spécifiques" not in r.stdout:
    echec("remplir: sommaire", "le sommaire du PDF ne reflète pas le contenu réel "
          f"(phase 4 de la fixture absente de la page 2) :\n{r.stdout[:300]}")
if "Modules Spécifiques" in r.stdout:
    echec("remplir: sommaire", "le sommaire du PDF contient encore le cache périmé d'origine")
if re.search(r"Sommaire ?\.{3,}", r.stdout):
    echec("remplir: sommaire", "le sommaire se liste lui-même")
ok("remplir: sommaire réel — champs rafraîchis, phases présentes, pas d'auto-entrée")

r = subprocess.run([sys.executable, str(REMPLIR), str(gabarit), str(ctx_file),
                    str(livrable)], capture_output=True, text=True)
if r.returncode == 0 or "écras" not in r.stdout.lower():
    echec("remplir: écrasement", f"code {r.returncode}, sortie : {r.stdout!r}")
ok("remplir: sortie existante → refus, jamais d'écrasement")

ctx_file2 = tmp / "ctx-ampute2.json"
ctx_file2.write_text(json.dumps(ctx_ampute))
sortie2 = tmp / "livrable2.docx"
r = subprocess.run([sys.executable, str(REMPLIR), str(gabarit), str(ctx_file2),
                    str(sortie2)], capture_output=True, text=True)
if r.returncode == 0 or "date" not in r.stdout or sortie2.exists():
    echec("remplir: contexte incomplet", f"code {r.returncode}, fichier créé : {sortie2.exists()}")
ok("remplir: variable manquante → refus avant tout rendu, rien n'est écrit")

# ── 17. Vérifier le livrable : chiffres énumérés, résidus bloquants ───────
VERIF = RACINE / "instruments" / "verif_livrable.py"
r = subprocess.run([sys.executable, str(VERIF), str(livrable)],
                   capture_output=True, text=True)
if r.returncode != 0:
    echec("verif", f"code {r.returncode}\n{r.stdout}{r.stderr}")
if "12 500 000" not in r.stdout or not re.search(r"\d+ chiffres? à sourcer", r.stdout):
    echec("verif: chiffres", f"énumération absente :\n{r.stdout[:400]}")
if "typographie" not in r.stdout.lower():
    echec("verif: typographie", "aucun signalement typographique dans la sortie")
ok("verif: chiffres énumérés avec décompte, typographie signalée")

troue = tmp / "livrable-troue.docx"
d = docx.Document()
d.add_paragraph("Le montant est de {{ montant }} FCFA.")
d.save(str(troue))
r = subprocess.run([sys.executable, str(VERIF), str(troue)],
                   capture_output=True, text=True)
if r.returncode != 1 or "montant" not in r.stdout:
    echec("verif: résidu", f"code {r.returncode}\n{r.stdout}")
ok("verif: résidu de gabarit dans un livrable → code 1")

vide2 = tmp / "livrable-vide.docx"
docx.Document().save(str(vide2))
r = subprocess.run([sys.executable, str(VERIF), str(vide2)],
                   capture_output=True, text=True)
if r.returncode != 2:
    echec("verif: refus du vide", f"code {r.returncode}")
ok("verif: livrable sans texte → code 2, jamais un OK muet")

# ── 18. doctor : la machine sait dire « je ne peux pas fonctionner ici » ──
import os
DOCTOR = RACINE / "instruments" / "doctor.py"
r = subprocess.run([sys.executable, str(DOCTOR)], capture_output=True, text=True)
if r.returncode != 0:
    echec("doctor", f"code {r.returncode} sur une machine saine\n{r.stdout}{r.stderr}")
for requis in ["docxtpl", "docxcompose", "libreoffice", "pdfinfo", "templates"]:
    if not re.search(rf"OK\s+.*{requis}", r.stdout):
        echec("doctor", f"requis non listé OK : {requis}\n{r.stdout}")
ok("doctor: machine saine → code 0, chaque requis listé OK")

env_ampute = dict(os.environ, PATH="/nonexistent")
r = subprocess.run([sys.executable, str(DOCTOR)], capture_output=True, text=True,
                   env=env_ampute)
if r.returncode == 0 or not re.search(r"ABSENT\s+.*libreoffice", r.stdout) \
        or "apt install" not in r.stdout:
    echec("doctor: binaire absent", f"code {r.returncode}\n{r.stdout}")
ok("doctor: binaire requis absent → code 1, commande d'installation donnée")

r = subprocess.run([sys.executable, str(DOCTOR), "--racine", str(tmp)],
                   capture_output=True, text=True)
if r.returncode == 0 or "templates" not in r.stdout:
    echec("doctor: templates vides", f"code {r.returncode}\n{r.stdout}")
ok("doctor: templates/ vide → code 1, la machine refuse de passer pour installée")

# ── 19. Câblage : skill mince + installation idempotente au bon profil ────
SKILL = RACINE / "skill" / "SKILL.md"
if not SKILL.exists():
    echec("skill", "skill/SKILL.md absent")
skill_txt = SKILL.read_text()
fm = re.match(r"---\n(.*?)\n---\n", skill_txt, re.S)
if not fm or "name: machine-redaction" not in fm.group(1) \
        or "description:" not in fm.group(1):
    echec("skill", "frontmatter sans name/description")
if len(fm.group(1).splitlines()) > 10:
    echec("skill", f"frontmatter obèse ({len(fm.group(1).splitlines())} lignes) — "
          "leçon R6 : coût de contexte permanent")
refs = set(re.findall(r"instruments/([a-z_]+\.py)", skill_txt))
if not refs:
    echec("skill", "le corps ne référence aucun instrument")
for ref in refs:
    if not (RACINE / "instruments" / ref).exists():
        echec("skill", f"référence morte (leçon R11) : instruments/{ref}")
ok(f"skill: frontmatter mince, {len(refs)} instruments référencés, zéro référence morte")

faux_home = tmp / "home"
faux_home.mkdir()
import site
env_home = dict(os.environ, HOME=str(faux_home),
                PYTHONPATH=site.getusersitepackages())   # les modules pip --user
                                                         # suivent HOME : on les garde
for passage in (1, 2):    # deux passages : l'installation doit être idempotente
    r = subprocess.run(["bash", str(RACINE / "install.sh")],
                       capture_output=True, text=True, env=env_home)
    if r.returncode != 0:
        echec("install", f"passage {passage} : code {r.returncode}\n{r.stdout}{r.stderr}")
lien = faux_home / ".claude-mecid" / "skills" / "machine-redaction"
if not lien.is_symlink() or not (lien / "SKILL.md").read_text() == skill_txt:
    echec("install", "lien absent du bon profil ou contenu divergent")
ok("install: lien symbolique au bon profil (~/.claude-mecid/skills), idempotent")

# ── 21. Toilettage typographique : espaces doubles, insécables, guillemets ─
TOILETTE = RACINE / "instruments" / "toiletter_gabarit.py"
sale = tmp / "typo-sale.docx"
d = docx.Document()
d.add_paragraph("Un montant de  22 000 000  (vingt-deux millions) FCFA.")
d.add_paragraph("permettant de : rien")
d.add_paragraph("la transformation digitale; elles la mènent.")
d.add_paragraph("« citation »")
d.save(str(sale))
r = subprocess.run([sys.executable, str(TOILETTE), str(sale), str(tmp / "typo-propre.docx")],
                   capture_output=True, text=True)
if r.returncode != 0:
    echec("toilettage", f"code {r.returncode}\n{r.stdout}{r.stderr}")
txt_t = textes_docx(tmp / "typo-propre.docx")["word/document.xml"]
attendus_typo = [
    "de 22 000 000 (vingt-deux millions)",
    "permettant de : rien",
    "digitale ; elles",
    "« citation »",
]
for att in attendus_typo:
    if att not in txt_t:
        echec("toilettage", f"correction absente : {att!r} dans {txt_t!r}")
ok("toilettage: doubles espaces, insécables et guillemets corrigés")

# ── 22. Finitions du gabarit PTF : alignements, marges, solidarité ────────
xml_gab = zipfile.ZipFile(gabarit).read("word/document.xml").decode("utf-8")
i = xml_gab.find("{{ item.cost }}")
pPr_cout = xml_gab[xml_gab.rfind("<w:p ", 0, i):i]
if 'w:jc w:val="right"' not in pPr_cout:
    echec("finitions", "montants du tableau financier non alignés à droite")
i = xml_gab.find("{{ item.count }}")
pPr_count = xml_gab[xml_gab.rfind("<w:p ", 0, i):i]
if 'w:jc w:val="center"' not in pPr_count:
    echec("finitions", "effectifs du tableau équipe non centrés")
m = re.search(r'<w:pgMar[^/]*w:footer="(\d+)"[^/]*/>|<w:pgMar[^/]*/>', xml_gab)
footer_twips = int(re.search(r'w:footer="(\d+)"', m.group(0)).group(1))
if footer_twips < 700:
    echec("finitions", f"zone de pied de page trop courte ({footer_twips} twips) — "
          "le corps déborde dans le footer")
i = xml_gab.find("{{ item.solution_name }}")
if i == -1:
    i = xml_gab.find("{{ solution_name }}")
pPr_sol = xml_gab[xml_gab.rfind("<w:p ", 0, i):i]
if "w:keepNext" not in pPr_sol:
    echec("finitions", "titre de solution orphelin possible (keepNext absent)")
m = re.search(r'<wp:positionH relativeFrom="column">\s*<wp:posOffset>(\d+)</wp:posOffset>', xml_gab)
if m and int(m.group(1)) > 2000000:
    echec("finitions", f"cachet ancré à {m.group(1)} EMU — il chevauche la "
          "colonne signature du client au lieu de couvrir celle de GINUTECH")
ok("finitions: montants à droite, effectifs centrés, footer élargi, cachet côté GINUTECH")

# ── 23. Contrat de forme : police unique, palette et échelle fermées ──────
forme = json.loads((RACINE / "forme.json").read_text())
r = subprocess.run(["pdffonts", str(pdf_livrable)], capture_output=True, text=True)
familles = {l.split()[0].split("+")[-1] for l in r.stdout.splitlines()[2:] if l.strip()}
etrangeres = {f for f in familles if forme["police"] not in f}
if etrangeres:
    echec("forme: polices", f"familles hors contrat dans le PDF : {sorted(etrangeres)} "
          f"(attendu : {forme['police']} seule)")
ok(f"forme: police unique dans le PDF ({sorted(familles)})")

xml_liv = zipfile.ZipFile(livrable).read("word/document.xml").decode("utf-8")
fills = set(re.findall(r'w:fill="([0-9A-Fa-f]{6}|auto)"', xml_liv))
hors_palette = fills - set(forme["couleurs_autorisees"])
if hors_palette:
    echec("forme: palette", f"fonds hors contrat : {sorted(hors_palette)}")
couleurs_txt = set(re.findall(r'<w:color w:val="([0-9A-Fa-f]{6})"/>', xml_liv))
hors_palette = couleurs_txt - set(forme["couleurs_autorisees"])
if hors_palette:
    echec("forme: palette texte", f"couleurs hors contrat : {sorted(hors_palette)}")
tailles = {int(t) for t in re.findall(r'<w:sz w:val="(\d+)"/>', xml_liv)}
hors_echelle = tailles - set(forme["tailles_autorisees"])
if hors_echelle:
    echec("forme: échelle", f"tailles hors contrat (demi-points) : {sorted(hors_echelle)}")
ok("forme: palette et échelle typographique fermées")

corps_liv = textes_docx(livrable)["word/document.xml"]
if "Durée (en semaines)" not in corps_liv or "— Usage interne" not in corps_liv:
    echec("forme: micro-textes", "'en semaines' ou tiret cadratin absent")
r = subprocess.run(["pdfinfo", str(pdf_livrable)], capture_output=True, text=True)
titre_pdf = next((l.split(":", 1)[1].strip() for l in r.stdout.splitlines()
                  if l.startswith("Title")), "")
if not titre_pdf or titre_pdf == "Word Document":
    echec("forme: métadonnées", f"titre PDF : {titre_pdf!r}")
ok(f"forme: micro-textes corrigés, métadonnées PDF renseignées ({titre_pdf!r})")

# ── 20. AF-xxx : chaque valeur chiffrée de la base de faits a une source ──
AF = RACINE / "instruments" / "af_contexte.py"
r = subprocess.run([sys.executable, str(AF), str(ctx_file)],
                   capture_output=True, text=True)
if r.returncode != 0:
    echec("af", f"code {r.returncode} sur un contexte sourcé\n{r.stdout}{r.stderr}")
if "AF-004" not in r.stdout or "somme des coûts" not in r.stdout:
    echec("af: registre", f"registre AF absent de la sortie :\n{r.stdout[:300]}")
ok("af: contexte sourcé → code 0, registre AF imprimé")

ctx_sans_af = {k: v for k, v in contexte.items() if k != "_af"}
f_sans_af = tmp / "ctx-sans-af.json"
f_sans_af.write_text(json.dumps(ctx_sans_af))
r = subprocess.run([sys.executable, str(AF), str(f_sans_af)],
                   capture_output=True, text=True)
if r.returncode != 1 or "employee_count" not in r.stdout:
    echec("af: delta", f"code {r.returncode}, valeur non sourcée non nommée :\n{r.stdout[:300]}")
ok("af: valeur chiffrée sans source → code 1, chemin nommé")

ctx_af_morte = dict(contexte)
ctx_af_morte["_af"] = contexte["_af"] + [
    {"id": "AF-099", "chemin": "chemin.inexistant", "source": "rien"}]
f_morte = tmp / "ctx-af-morte.json"
f_morte.write_text(json.dumps(ctx_af_morte))
r = subprocess.run([sys.executable, str(AF), str(f_morte)],
                   capture_output=True, text=True)
if r.returncode != 2 or "AF-099" not in r.stdout:
    echec("af: entrée morte", f"code {r.returncode}\n{r.stdout[:300]}")
ok("af: entrée AF qui ne couvre rien → code 2, nommée")

sortie_af = tmp / "livrable-sans-af.docx"
r = subprocess.run([sys.executable, str(REMPLIR), str(gabarit), str(f_sans_af),
                    str(sortie_af)], capture_output=True, text=True)
if r.returncode == 0 or "AF" not in r.stdout or sortie_af.exists():
    echec("remplir: base de faits non sourcée",
          f"code {r.returncode}, fichier créé : {sortie_af.exists()}\n{r.stdout[:300]}")
ok("remplir: base de faits non sourcée → refus, rien n'est écrit")

# ── 24. Pont spec-kit → CDC : extraction et fermeture des exigences ───────
PONT = RACINE / "instruments" / "spec_vers_contexte.py"
SPEC = ICI / "fixtures" / "spec_kmerfoot.md"
ctx_cdc_file = tmp / "cdc-depuis-spec.json"
r = subprocess.run([sys.executable, str(PONT), str(SPEC), str(ctx_cdc_file)],
                   capture_output=True, text=True)
if r.returncode != 0:
    echec("pont: extraction", f"code {r.returncode}\n{r.stdout}{r.stderr}")
amorce = json.loads(ctx_cdc_file.read_text())
if len(amorce["_spec"]["exigences"]) != 6:
    echec("pont: extraction", f"{len(amorce['_spec']['exigences'])} exigences au lieu de 6")
if amorce["_spec"]["exigences"][0]["id"] != "FR-001" \
        or "ligne" not in amorce["_spec"]["exigences"][0]:
    echec("pont: extraction", "exigence sans id ou sans ligne source")
if "paris" not in amorce.get("project_name", "").lower():
    echec("pont: extraction", f"project_name non repris du titre : {amorce.get('project_name')!r}")
ok("pont: 6 exigences extraites avec ligne source, contexte amorcé")

vide_spec = tmp / "sans-fr.md"
vide_spec.write_text("# Feature Specification: rien\n\nDu texte sans exigence.\n")
r = subprocess.run([sys.executable, str(PONT), str(vide_spec), str(tmp / "x.json")],
                   capture_output=True, text=True)
if r.returncode != 2:
    echec("pont: refus du vide", f"code {r.returncode}")
ok("pont: spec sans exigence → code 2")

ctx_complet = dict(amorce)
ctx_complet["functional_modules"] = [{
    "module_name": "Suivi et paris",
    "submodules": [{"submodule_name": "Suivi", "features": [
        {"feature_name": "Scores et classements (FR-001)"},
        {"feature_name": "Prise de paris (FR-002, FR-003)"},
        {"feature_name": "Portefeuille Mobile Money (FR-004, FR-005)"},
        {"feature_name": "Back-office (FR-006)"}]}]}]
f_complet = tmp / "cdc-complet.json"
f_complet.write_text(json.dumps(ctx_complet))
r = subprocess.run([sys.executable, str(PONT), str(SPEC), "--verifier", str(f_complet)],
                   capture_output=True, text=True)
if r.returncode != 0:
    echec("pont: fermeture", f"code {r.returncode} sur un CDC complet\n{r.stdout}")
ok("pont: fermeture vérifiée — chaque FR de la spec est dans le CDC")

ctx_troue = json.loads(f_complet.read_text())
ctx_troue["functional_modules"][0]["submodules"][0]["features"][2]["feature_name"] = \
    "Portefeuille Mobile Money (FR-005)"   # FR-004 disparaît
f_troue = tmp / "cdc-troue.json"
f_troue.write_text(json.dumps(ctx_troue))
r = subprocess.run([sys.executable, str(PONT), str(SPEC), "--verifier", str(f_troue)],
                   capture_output=True, text=True)
if r.returncode != 1 or "FR-004" not in r.stdout:
    echec("pont: delta", f"code {r.returncode}, FR-004 non nommée\n{r.stdout}")
ok("pont: exigence absente du CDC → code 1, FR nommée (delta)")

ctx_invente = json.loads(f_complet.read_text())
ctx_invente["functional_modules"][0]["submodules"][0]["features"][0]["feature_name"] += " (FR-099)"
f_invente = tmp / "cdc-invente.json"
f_invente.write_text(json.dumps(ctx_invente))
r = subprocess.run([sys.executable, str(PONT), str(SPEC), "--verifier", str(f_invente)],
                   capture_output=True, text=True)
if r.returncode != 2 or "FR-099" not in r.stdout:
    echec("pont: invention", f"code {r.returncode}, FR-099 non nommée\n{r.stdout}")
ok("pont: FR citée mais absente de la spec → code 2 (invention)")

# ── 14. Le rendu se convertit en PDF (LibreOffice) ─────────────────────────
r = subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf",
                    "--outdir", str(tmp), str(rendu)],
                   capture_output=True, text=True, timeout=120)
pdf = tmp / "ptf-rendu.pdf"
if r.returncode != 0 or not pdf.exists() or pdf.stat().st_size < 10000:
    echec("pdf", f"code {r.returncode}, pdf={'présent' if pdf.exists() else 'absent'}")
ok("pdf: conversion LibreOffice aboutie")

if verifs == 0:
    echec("porte", "zéro vérification exécutée")
print(f"\nPORTE VERTE — {verifs} vérifications")
shutil.rmtree(tmp)
